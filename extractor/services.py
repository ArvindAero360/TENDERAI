"""
extractor/services.py
─────────────────────
All extraction logic ported from the Flask app.py.
Pure Python — no Django imports — so it can be tested independently.
"""
import re as _re
import time
import json

import requests as req_lib
import pdfplumber
import docx as _docx

# ── CID cleaner ──────────────────────────────────────────────────────────────
_CID_PAT = _re.compile(r'\(cid:\d+\)')

def _clean_cid(text):
    if '(cid:' not in text:
        return text
    cleaned = _CID_PAT.sub(' ', text)
    cleaned = _re.sub(r'[^\S\n]{2,}', ' ', cleaned)
    return cleaned


# ── Indian state data ────────────────────────────────────────────────────────
INDIA_STATE_NAMES = {
    'andhra pradesh', 'arunachal pradesh', 'assam', 'bihar', 'chhattisgarh',
    'goa', 'gujarat', 'haryana', 'himachal pradesh', 'jharkhand', 'karnataka',
    'kerala', 'madhya pradesh', 'maharashtra', 'manipur', 'meghalaya',
    'mizoram', 'nagaland', 'odisha', 'punjab', 'rajasthan', 'sikkim',
    'tamil nadu', 'telangana', 'tripura', 'uttar pradesh', 'uttarakhand',
    'west bengal', 'delhi', 'jammu & kashmir', 'ladakh', 'puducherry',
    'chandigarh', 'andaman & nicobar islands', 'lakshadweep',
    'dadra and nagar haveli', 'daman and diu',
}

_STATE_DISPLAY = {s: s.title() for s in INDIA_STATE_NAMES}
_STATE_DISPLAY.update({
    'jammu & kashmir': 'Jammu & Kashmir',
    'andaman & nicobar islands': 'Andaman & Nicobar Islands',
    'dadra and nagar haveli': 'Dadra and Nagar Haveli',
    'daman and diu': 'Daman and Diu',
})

_STATE_ABBREV = {
    'mh': 'Maharashtra', 'ka': 'Karnataka', 'tn': 'Tamil Nadu', 'gj': 'Gujarat',
    'rj': 'Rajasthan', 'up': 'Uttar Pradesh', 'wb': 'West Bengal', 'ts': 'Telangana',
    'ap': 'Andhra Pradesh', 'kl': 'Kerala', 'pb': 'Punjab', 'hr': 'Haryana',
    'mp': 'Madhya Pradesh', 'br': 'Bihar', 'or': 'Odisha', 'od': 'Odisha',
    'as': 'Assam', 'jh': 'Jharkhand', 'cg': 'Chhattisgarh', 'uk': 'Uttarakhand',
    'hp': 'Himachal Pradesh', 'ga': 'Goa', 'jk': 'Jammu & Kashmir',
    'j&k': 'Jammu & Kashmir', 'la': 'Ladakh', 'mn': 'Manipur',
    'ml': 'Meghalaya', 'tr': 'Tripura', 'nl': 'Nagaland', 'mz': 'Mizoram',
    'sk': 'Sikkim', 'ar': 'Arunachal Pradesh', 'py': 'Puducherry',
    'ch': 'Chandigarh', 'dl': 'Delhi', 'an': 'Andaman & Nicobar Islands',
    'ld': 'Lakshadweep',
}


def _normalize_state(raw):
    if not raw:
        return None
    s = raw.strip()
    abbrev_key = s.lower().replace('.', '').replace(' ', '')
    if abbrev_key in _STATE_ABBREV:
        return _STATE_ABBREV[abbrev_key]
    low = s.lower()
    if low in INDIA_STATE_NAMES:
        return _STATE_DISPLAY.get(low, s.title())
    for sname in sorted(INDIA_STATE_NAMES, key=len, reverse=True):
        if sname in low:
            return _STATE_DISPLAY.get(sname, sname.title())
    return s.title() if s else None


INDIA_CITY_STATE = {
    'mumbai': 'Maharashtra', 'pune': 'Maharashtra', 'nagpur': 'Maharashtra',
    'nashik': 'Maharashtra', 'aurangabad': 'Maharashtra', 'solapur': 'Maharashtra',
    'thane': 'Maharashtra', 'kolhapur': 'Maharashtra', 'navi mumbai': 'Maharashtra',
    'amravati': 'Maharashtra', 'nanded': 'Maharashtra', 'satara': 'Maharashtra',
    'latur': 'Maharashtra', 'sangli': 'Maharashtra', 'jalgaon': 'Maharashtra',
    'akola': 'Maharashtra', 'dhule': 'Maharashtra', 'ahmednagar': 'Maharashtra',
    'ratnagiri': 'Maharashtra', 'sindhudurg': 'Maharashtra', 'raigad': 'Maharashtra',
    'delhi': 'Delhi', 'new delhi': 'Delhi',
    'bangalore': 'Karnataka', 'bengaluru': 'Karnataka', 'mysore': 'Karnataka',
    'mysuru': 'Karnataka', 'hubli': 'Karnataka', 'dharwad': 'Karnataka',
    'mangalore': 'Karnataka', 'mangaluru': 'Karnataka', 'belagavi': 'Karnataka',
    'belgaum': 'Karnataka', 'kalaburagi': 'Karnataka', 'gulbarga': 'Karnataka',
    'davanagere': 'Karnataka', 'shimoga': 'Karnataka', 'shivamogga': 'Karnataka',
    'tumkur': 'Karnataka', 'bellary': 'Karnataka', 'ballari': 'Karnataka',
    'bidar': 'Karnataka', 'hassan': 'Karnataka', 'udupi': 'Karnataka',
    'chennai': 'Tamil Nadu', 'coimbatore': 'Tamil Nadu', 'madurai': 'Tamil Nadu',
    'trichy': 'Tamil Nadu', 'tiruchirappalli': 'Tamil Nadu', 'salem': 'Tamil Nadu',
    'tirunelveli': 'Tamil Nadu', 'vellore': 'Tamil Nadu', 'erode': 'Tamil Nadu',
    'tirupur': 'Tamil Nadu', 'thoothukudi': 'Tamil Nadu', 'thanjavur': 'Tamil Nadu',
    'dindigul': 'Tamil Nadu', 'kanchipuram': 'Tamil Nadu', 'cuddalore': 'Tamil Nadu',
    'ahmedabad': 'Gujarat', 'surat': 'Gujarat', 'vadodara': 'Gujarat',
    'rajkot': 'Gujarat', 'gandhinagar': 'Gujarat', 'bhavnagar': 'Gujarat',
    'jamnagar': 'Gujarat', 'junagadh': 'Gujarat', 'anand': 'Gujarat',
    'navsari': 'Gujarat', 'morbi': 'Gujarat', 'mehsana': 'Gujarat',
    'jaipur': 'Rajasthan', 'jodhpur': 'Rajasthan', 'udaipur': 'Rajasthan',
    'kota': 'Rajasthan', 'bikaner': 'Rajasthan', 'ajmer': 'Rajasthan',
    'alwar': 'Rajasthan', 'bharatpur': 'Rajasthan', 'sikar': 'Rajasthan',
    'pali': 'Rajasthan', 'sri ganganagar': 'Rajasthan', 'barmer': 'Rajasthan',
    'chittorgarh': 'Rajasthan', 'bhilwara': 'Rajasthan', 'jaisalmer': 'Rajasthan',
    'lucknow': 'Uttar Pradesh', 'kanpur': 'Uttar Pradesh', 'agra': 'Uttar Pradesh',
    'varanasi': 'Uttar Pradesh', 'allahabad': 'Uttar Pradesh', 'prayagraj': 'Uttar Pradesh',
    'meerut': 'Uttar Pradesh', 'noida': 'Uttar Pradesh', 'ghaziabad': 'Uttar Pradesh',
    'mathura': 'Uttar Pradesh', 'aligarh': 'Uttar Pradesh', 'bareilly': 'Uttar Pradesh',
    'moradabad': 'Uttar Pradesh', 'gorakhpur': 'Uttar Pradesh', 'firozabad': 'Uttar Pradesh',
    'saharanpur': 'Uttar Pradesh', 'muzaffarnagar': 'Uttar Pradesh', 'rampur': 'Uttar Pradesh',
    'jhansi': 'Uttar Pradesh', 'faizabad': 'Uttar Pradesh', 'ayodhya': 'Uttar Pradesh',
    'kolkata': 'West Bengal', 'howrah': 'West Bengal', 'durgapur': 'West Bengal',
    'asansol': 'West Bengal', 'siliguri': 'West Bengal', 'bardhaman': 'West Bengal',
    'burdwan': 'West Bengal', 'malda': 'West Bengal', 'baharampur': 'West Bengal',
    'jalpaiguri': 'West Bengal', 'kharagpur': 'West Bengal',
    'hyderabad': 'Telangana', 'warangal': 'Telangana', 'karimnagar': 'Telangana',
    'secunderabad': 'Telangana', 'nizamabad': 'Telangana', 'khammam': 'Telangana',
    'mahbubnagar': 'Telangana', 'nalgonda': 'Telangana', 'adilabad': 'Telangana',
    'visakhapatnam': 'Andhra Pradesh', 'vijayawada': 'Andhra Pradesh',
    'guntur': 'Andhra Pradesh', 'vizag': 'Andhra Pradesh', 'tirupati': 'Andhra Pradesh',
    'amaravati': 'Andhra Pradesh', 'kakinada': 'Andhra Pradesh', 'nellore': 'Andhra Pradesh',
    'kurnool': 'Andhra Pradesh', 'rajahmundry': 'Andhra Pradesh', 'eluru': 'Andhra Pradesh',
    'kadapa': 'Andhra Pradesh', 'anantapur': 'Andhra Pradesh',
    'thiruvananthapuram': 'Kerala', 'kochi': 'Kerala', 'kozhikode': 'Kerala',
    'thrissur': 'Kerala', 'kollam': 'Kerala', 'trivandrum': 'Kerala',
    'palakkad': 'Kerala', 'alappuzha': 'Kerala', 'kannur': 'Kerala',
    'malappuram': 'Kerala', 'ernakulam': 'Kerala', 'idukki': 'Kerala',
    'chandigarh': 'Chandigarh', 'amritsar': 'Punjab', 'ludhiana': 'Punjab',
    'jalandhar': 'Punjab', 'patiala': 'Punjab', 'bathinda': 'Punjab',
    'moga': 'Punjab', 'pathankot': 'Punjab', 'hoshiarpur': 'Punjab',
    'gurugram': 'Haryana', 'gurgaon': 'Haryana', 'faridabad': 'Haryana',
    'ambala': 'Haryana', 'hisar': 'Haryana', 'rohtak': 'Haryana',
    'karnal': 'Haryana', 'panipat': 'Haryana', 'sonipat': 'Haryana',
    'bhiwani': 'Haryana', 'sirsa': 'Haryana', 'yamunanagar': 'Haryana',
    'bhopal': 'Madhya Pradesh', 'indore': 'Madhya Pradesh', 'gwalior': 'Madhya Pradesh',
    'jabalpur': 'Madhya Pradesh', 'ujjain': 'Madhya Pradesh', 'sagar': 'Madhya Pradesh',
    'dewas': 'Madhya Pradesh', 'satna': 'Madhya Pradesh', 'ratlam': 'Madhya Pradesh',
    'rewa': 'Madhya Pradesh', 'murwara': 'Madhya Pradesh', 'singrauli': 'Madhya Pradesh',
    'patna': 'Bihar', 'gaya': 'Bihar', 'muzaffarpur': 'Bihar',
    'bhagalpur': 'Bihar', 'darbhanga': 'Bihar', 'purnia': 'Bihar',
    'arrah': 'Bihar', 'begusarai': 'Bihar', 'katihar': 'Bihar',
    'bhubaneswar': 'Odisha', 'cuttack': 'Odisha', 'rourkela': 'Odisha',
    'brahmapur': 'Odisha', 'sambalpur': 'Odisha', 'puri': 'Odisha',
    'berhampur': 'Odisha', 'balasore': 'Odisha',
    'guwahati': 'Assam', 'dibrugarh': 'Assam', 'silchar': 'Assam',
    'jorhat': 'Assam', 'nagaon': 'Assam', 'tinsukia': 'Assam',
    'ranchi': 'Jharkhand', 'jamshedpur': 'Jharkhand', 'dhanbad': 'Jharkhand',
    'bokaro': 'Jharkhand', 'deoghar': 'Jharkhand', 'hazaribagh': 'Jharkhand',
    'raipur': 'Chhattisgarh', 'bilaspur': 'Chhattisgarh', 'durg': 'Chhattisgarh',
    'korba': 'Chhattisgarh', 'rajnandgaon': 'Chhattisgarh',
    'dehradun': 'Uttarakhand', 'haridwar': 'Uttarakhand', 'roorkee': 'Uttarakhand',
    'haldwani': 'Uttarakhand', 'rudrapur': 'Uttarakhand', 'kashipur': 'Uttarakhand',
    'shimla': 'Himachal Pradesh', 'manali': 'Himachal Pradesh',
    'dharamshala': 'Himachal Pradesh', 'solan': 'Himachal Pradesh',
    'mandi': 'Himachal Pradesh', 'kullu': 'Himachal Pradesh',
    'panaji': 'Goa', 'margao': 'Goa', 'vasco': 'Goa',
    'mapusa': 'Goa', 'ponda': 'Goa',
    'srinagar': 'Jammu & Kashmir', 'jammu': 'Jammu & Kashmir',
    'anantnag': 'Jammu & Kashmir', 'baramulla': 'Jammu & Kashmir',
    'leh': 'Ladakh', 'kargil': 'Ladakh',
    'imphal': 'Manipur', 'shillong': 'Meghalaya', 'agartala': 'Tripura',
    'kohima': 'Nagaland', 'dimapur': 'Nagaland', 'aizawl': 'Mizoram',
    'gangtok': 'Sikkim', 'itanagar': 'Arunachal Pradesh',
    'puducherry': 'Puducherry', 'pondicherry': 'Puducherry',
    'port blair': 'Andaman & Nicobar Islands', 'kavaratti': 'Lakshadweep',
    'ranga reddy': 'Telangana', 'medchal': 'Telangana',
}

OSM_NOMINATIM_URL = 'https://nominatim.openstreetmap.org/search'
OSM_NOMINATIM_HEADERS = {
    'User-Agent': 'TenderIQ/4.1 (indian-tender-extraction; opensource)',
    'Accept-Language': 'en',
}
_geocode_cache: dict = {}


def _state_from_osm_address(address: dict):
    if not address:
        return None
    for key in ('state', 'state_district', 'county', 'region'):
        val = address.get(key, '').strip().lower()
        if not val:
            continue
        if val in INDIA_STATE_NAMES:
            return _STATE_DISPLAY.get(val, val.title())
        for sname in sorted(INDIA_STATE_NAMES, key=len, reverse=True):
            if sname in val or val in sname:
                return _STATE_DISPLAY.get(sname, sname.title())
    return None


def _geocode_state_via_osm(location_str: str):
    cache_key = location_str.lower().strip()
    if cache_key in _geocode_cache:
        return _geocode_cache[cache_key]
    try:
        resp = req_lib.get(
            OSM_NOMINATIM_URL,
            params={'q': location_str, 'countrycodes': 'in', 'format': 'json',
                    'addressdetails': 1, 'limit': 5},
            headers=OSM_NOMINATIM_HEADERS,
            timeout=8,
        )
        resp.raise_for_status()
        results = resp.json()
        time.sleep(1.1)
        if not results:
            _geocode_cache[cache_key] = None
            return None
        for result in results:
            state = _state_from_osm_address(result.get('address', {}))
            if state:
                _geocode_cache[cache_key] = state
                return state
        _geocode_cache[cache_key] = None
        return None
    except Exception as e:
        print(f'[osm-geocode] Error for {location_str!r}: {e}')
        return None


def _fallback_state_from_dict(location_str: str):
    loc = location_str.lower().strip()
    for sname in INDIA_STATE_NAMES:
        if sname in loc:
            return _STATE_DISPLAY.get(sname, sname.title())
    best_key, best_state = '', None
    for city, state in INDIA_CITY_STATE.items():
        if city in loc and len(city) > len(best_key):
            best_key, best_state = city, state
    if best_state:
        return best_state
    tokens = _re.split(r'[,\s/\-]+', loc)
    for token in tokens:
        token = token.strip()
        if len(token) < 3:
            continue
        if token in INDIA_CITY_STATE:
            return INDIA_CITY_STATE[token]
    return None


def get_state_from_location(location_str: str):
    if not location_str:
        return None
    quick = _fallback_state_from_dict(location_str)
    if quick:
        return quick
    if _re.search(
            r'\b(DISTT?|PIN[-:\s]*\d{6})\b'
            r'|(?<![A-Z])(CG|UP|MP|MH|KA|TN|AP|TS|RJ|GJ|PB|HR|HP'
            r'|UK|JH|OD|AS|BR|WB|JK|DL)(?![A-Z])',
            location_str.upper()):
        _, addr_state = parse_location_from_address(location_str)
        if addr_state:
            return addr_state
    return _geocode_state_via_osm(location_str)


def extract_state_from_doc(doc_text: str):
    if not doc_text:
        return None
    text_lower = doc_text.lower()
    for sname in sorted(INDIA_STATE_NAMES, key=len, reverse=True):
        if sname in text_lower:
            return _STATE_DISPLAY.get(sname, sname.title())
    return None


ABBREV_TO_STATE = {
    'CG': 'Chhattisgarh', 'UP': 'Uttar Pradesh', 'MP': 'Madhya Pradesh',
    'MH': 'Maharashtra', 'KA': 'Karnataka', 'TN': 'Tamil Nadu',
    'AP': 'Andhra Pradesh', 'TS': 'Telangana', 'RJ': 'Rajasthan',
    'GJ': 'Gujarat', 'PB': 'Punjab', 'HR': 'Haryana',
    'HP': 'Himachal Pradesh', 'UK': 'Uttarakhand', 'JH': 'Jharkhand',
    'OD': 'Odisha', 'AS': 'Assam', 'BR': 'Bihar',
    'WB': 'West Bengal', 'JK': 'Jammu & Kashmir', 'J&K': 'Jammu & Kashmir',
    'DL': 'Delhi', 'GA': 'Goa', 'MN': 'Manipur',
    'ML': 'Meghalaya', 'TR': 'Tripura', 'NL': 'Nagaland',
    'MZ': 'Mizoram', 'SK': 'Sikkim', 'AR': 'Arunachal Pradesh',
    'PY': 'Puducherry', 'CH': 'Chandigarh', 'LA': 'Ladakh',
}

_ADDR_ORG_WORDS = frozenset([
    'BN', 'CRPF', 'BSF', 'CISF', 'SSB', 'ITBP', 'NSG', 'SIG', 'HQ', 'CAMP',
    'COMPLEX', 'NEAR', 'LINE', 'SECTOR', 'BLOCK', 'UNIT', 'POST', 'BASE', 'WING',
    'POLICE', 'WIRELESS', 'ARMY', 'NAVY', 'AIR', 'FORCE', 'BATTALION', 'COMPANY',
    'EW', 'SQN', 'RCT', 'CTC', 'GC', 'GREF', 'PWD', 'PWO', 'DISTT', 'DIST',
    'PIN', 'INDIA', 'THE', 'AND', 'FOR', 'WITH', 'NEAR', 'ROAD', 'STREET',
    'KASHMIR', 'HARYANA', 'PUNJAB', 'GUJARAT', 'RAJASTHAN', 'MAHARASHTRA',
    'KARNATAKA', 'TELANGANA', 'ANDHRA', 'PRADESH', 'BENGAL', 'ORISSA', 'ODISHA',
    'UTTARAKHAND', 'JHARKHAND', 'CHHATTISGARH', 'ASSAM', 'BIHAR', 'KERALA',
    'TRIPURA', 'MANIPUR', 'MEGHALAYA', 'NAGALAND', 'MIZORAM', 'SIKKIM', 'GOA',
])


def parse_location_from_address(addr: str) -> tuple:
    if not addr:
        return None, None
    addr_clean = _re.sub(r'\s+', ' ', addr.strip().upper())
    abbrev_state = None
    m_ab = _re.search(
        r'(?<![A-Z])(CG|UP|MP|MH|KA|TN|AP|TS|RJ|GJ|PB|HR|HP|UK|JH|OD|AS|BR'
        r'|WB|JK|DL|GA|MN|ML|TR|NL|MZ|SK|AR|PY|CH|LA|J&K)(?![A-Z])',
        addr_clean
    )
    if m_ab:
        abbrev_state = ABBREV_TO_STATE.get(m_ab.group(1))
    full_state = None
    addr_lower = addr_clean.lower()
    for sname in sorted(INDIA_STATE_NAMES, key=len, reverse=True):
        if sname in addr_lower:
            full_state = _STATE_DISPLAY.get(sname, sname.title())
            break
    if not full_state and 'kashmir' in addr_lower:
        full_state = 'Jammu & Kashmir'
    state = abbrev_state or full_state
    m_d = _re.search(r'\b([A-Z][A-Z]+(?:\s+[A-Z][A-Z]+)?)\s+DISTT?(?:[\s\-:])', addr_clean)
    if m_d:
        candidate_dist = m_d.group(1).title()
        before_str = addr_clean[:m_d.start()].strip()
        tokens_before = [t for t in _re.split(r'[\s,/\-]+', before_str)
                         if t and t.isalpha() and len(t) >= 3 and t not in _ADDR_ORG_WORDS]
        if tokens_before:
            city = tokens_before[-1].title()
            loc = (f'{city}, {candidate_dist}' if city.lower() != candidate_dist.lower()
                   else candidate_dist)
        else:
            loc = candidate_dist
        return loc, state
    tokens = [t for t in _re.split(r'[\s,/\-:]+', addr_clean) if t]
    for t in reversed(tokens):
        if (t.isalpha() and len(t) >= 4 and t not in _ADDR_ORG_WORDS
                and not any(sn in t.lower() for sn in ('pradesh', 'kashmir', 'bengal'))):
            return t.title(), state
    return None, state


# ── Value helpers ────────────────────────────────────────────────────────────
_STR_NULL = frozenset(['null', 'none', 'n/a', 'na', 'nil', '-', '--', 'not found',
                       'not available', 'not applicable', 'unknown', ''])
_BOOL_LIKE = frozenset(['yes', 'no', 'true', 'false', '1', '0', 'null', 'none',
                        'n/a', 'na', '-', 'nil'])


def _is_valid_place(v):
    if not v:
        return False
    v = v.strip()
    if len(v) < 3:
        return False
    if not _re.match(r'^[A-Za-z]', v):
        return False
    if v.lower() in _BOOL_LIKE:
        return False
    if v.strip().isdigit():
        return False
    return True


def _clean_location(v):
    s = str(v).strip() if v is not None else None
    if not s:
        return ''
    if s.lower().strip() in _BOOL_LIKE:
        return ''
    if s.strip().isdigit():
        return ''
    return s


def _str(v):
    if v is None:
        return None
    s = str(v).strip()
    return None if s.lower() in _STR_NULL else s


def _notnull(v, d=None):
    return _str(v) or d


def _yn(v, d='No'):
    if v is None:
        return d
    return 'Yes' if str(v).strip().lower() in ('yes', 'true', '1') else 'No'


def _bool(v):
    if isinstance(v, bool):
        return v
    return str(v).strip().lower() in ('yes', 'true', '1') if v is not None else False


def _int(v):
    try:
        return int(v) if v is not None else None
    except (TypeError, ValueError):
        return None


def _dec(v):
    if v is None:
        return None
    try:
        v2 = _re.sub(r'[^\d.]', '', str(v))
        return '{:.2f}'.format(float(v2))
    except (TypeError, ValueError):
        return None


def _convert_gem_date(raw):
    if not raw:
        return None
    raw = raw.strip()
    if _re.match(r'\d{4}-\d{2}-\d{2}', raw):
        if '+' not in raw and raw.count('T') == 1:
            return raw + '+05:30'
        return raw
    m = _re.match(r'(\d{2})-(\d{2})-(\d{4})\s+(\d{2}:\d{2}:\d{2})', raw)
    if m:
        return f'{m.group(3)}-{m.group(2)}-{m.group(1)}T{m.group(4)}+05:30'
    m = _re.match(r'(\d{2})-(\d{2})-(\d{4})', raw)
    if m:
        return f'{m.group(3)}-{m.group(2)}-{m.group(1)}T00:00:00+05:30'
    return raw


def _fix_date(raw):
    v = _convert_gem_date(raw or '')
    return v if v else None


def _date_only(v):
    if v is None:
        return None
    s = str(v).strip()
    if s.lower() in _STR_NULL:
        return None
    m = _re.match(r'(\d{4}-\d{2}-\d{2})', s)
    return m.group(1) if m else None


# ── EMD / ePBG extractors ────────────────────────────────────────────────────
def _extract_emd_amount(doc_text):
    patterns = [
        r'\[TABLE\][^\n]*(?:EMD\s+Amount|ईएमड[ीि])[^\n]*\|\s*(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
        r'\[TABLE\][^\n]*\bEMD\b\s*(?:\(Rs\.?\))?[^\n]*\|\s*(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
        r'\[TABLE\][^\n]*Bid\s+Security[^\n]*\|\s*(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
        r'EMD\s+Amount\s*[:\|/\-]?\s*(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
        r'Earnest\s+Money\s+Deposit[^₹\d]{0,25}(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
        r'\bEMD\s*[:\-]\s*(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
        r'Bid\s+Security\s+(?:Amount\s*)?[:\-]?\s*(?:₹\s*|Rs?\.?\s*)?(\d[\d,\.]*)',
    ]
    for pat in patterns:
        m = _re.search(pat, doc_text, _re.I)
        if m:
            raw_val = m.group(1).replace(',', '').strip().rstrip('/-')
            val = _dec(raw_val)
            if val and float(val) > 0:
                return val
    return None


def _extract_epbg_percentage(doc_text):
    patterns = [
        r'\[TABLE\][^\n]*(?:ePBG\s+Percentage|ईपीबीजी)[^\n]*\|\s*(\d+(?:\.\d+)?)',
        r'\[TABLE\][^\n]*ePBG\s*%?[^|\d]{0,8}\|\s*(\d+(?:\.\d+)?)',
        r'ePBG\s+Percentage\s*(?:\(%\))?\s*[:\|]?\s*(\d+(?:\.\d+)?)\s*%?',
        r'ePBG\s+(?:percentage|%)\s*(?:is\s*)?[=:\-]?\s*(\d+(?:\.\d+)?)\s*%?',
        r'EMD\s*@\s*(\d+(?:\.\d+)?)\s*%',
        r'EMD\s+(?:Percentage|%)\s*[=:\-]\s*(\d+(?:\.\d+)?)',
        r'Performance\s+(?:Security|Bank\s+Guarantee)\s*@\s*(\d+(?:\.\d+)?)\s*%',
        r'Security\s+Deposit\s*@\s*(\d+(?:\.\d+)?)\s*%',
        r'(\d+(?:\.\d+)?)\s*%\s+of\s+(?:estimated|contract|tender|bid)\s+(?:cost|value|amount)',
    ]
    for pat in patterns:
        m = _re.search(pat, doc_text, _re.I)
        if m:
            try:
                val = '{:.2f}'.format(float(m.group(1)))
                if 0 < float(val) <= 100:
                    return val
            except (ValueError, TypeError):
                continue
    return None


# ── Turnover extractor ───────────────────────────────────────────────────────
_TURNOVER_UNIT = r'(\d[\d,\.]*)\s*(Lakh|Lac|Crore|Cr)s?\s*(?:\(s\))?'

_TURNOVER_PATTERNS = [
    ('bidder', _re.compile(
        r'(?i)\[TABLE\][^\n]*(?:Minimum\s+Average\s+Annual\s+Turnover.*?bidder'
        r'|Bidder.*?Annual.*?Turnover)[^\n]*\|\s*' + _TURNOVER_UNIT, _re.DOTALL)),
    ('bidder', _re.compile(
        r'(?i)(?:Minimum\s+Average\s+Annual\s+Turnover.*?bidder'
        r'|Bidder.*?Annual.*?Turnover).{0,300}?' + _TURNOVER_UNIT, _re.DOTALL)),
    ('bidder', _re.compile(
        r'(?i)Annual\s+Turnover\s+of\s+the\s+bidder.{0,200}?' + _TURNOVER_UNIT, _re.DOTALL)),
    ('oem', _re.compile(
        r'(?i)\[TABLE\][^\n]*OEM\s+(?:Average\s+)?Turnover[^\n]*\|\s*' + _TURNOVER_UNIT, _re.DOTALL)),
    ('oem', _re.compile(
        r'(?i)OEM\s+(?:Average\s+)?Turnover.{0,300}?' + _TURNOVER_UNIT, _re.DOTALL)),
]


def _normalise_turnover_unit(num_str: str, unit_str: str) -> str:
    try:
        n = float(num_str.replace(',', ''))
        n_fmt = f'{n:g}'
    except ValueError:
        n_fmt = num_str.strip()
    unit_low = unit_str.lower()
    unit_display = 'Lakh (s)' if unit_low in ('lakh', 'lac', 'lakhs', 'lacs') else 'Crore (s)'
    return f'{n_fmt} {unit_display}'


def _extract_turnover_from_doc(doc_text: str) -> dict:
    results = {'bidder': None, 'oem': None}
    for field_key, pat in _TURNOVER_PATTERNS:
        if results[field_key]:
            continue
        m = pat.search(doc_text)
        if m:
            results[field_key] = _normalise_turnover_unit(m.group(1), m.group(2))
    return results


def _format_turnover_criteria(turnover: dict):
    bidder = turnover.get('bidder')
    oem    = turnover.get('oem')
    if bidder and oem:
        return f'Bidder: {bidder}, OEM: {oem}'
    if bidder:
        return bidder
    if oem:
        return f'OEM: {oem}'
    return None


# ── Pre-bid extractor ────────────────────────────────────────────────────────
def extract_prebid_full(doc_text):
    DT_PAT = _re.compile(r'(\d{2}-\d{2}-\d{4}\s+\d{2}:\d{2}:\d{2})')
    lines = doc_text.split('\n')
    for i, line in enumerate(lines):
        if not line.startswith('[TABLE]'):
            continue
        if not _re.search(r'(?i)Pre[\s\-]*Bid\s+Date', line):
            continue
        for data_line in lines[i + 1: i + 6]:
            data_line = data_line.strip()
            if not data_line:
                continue
            body = _re.sub(r'^\[TABLE\]\s*', '', data_line)
            m = DT_PAT.match(body)
            if m:
                dt = m.group(1).strip()
                rest = body[m.end():].strip().lstrip('|').strip()
                venue_parts = [rest] if rest else []
                for cont in lines[i + 2: i + 10]:
                    cont = cont.strip()
                    if not cont or cont.startswith('[TABLE]'):
                        break
                    venue_parts.append(cont)
                venue = ' '.join(p for p in venue_parts if p)
                return dt, venue
    header_idx = None
    for i, line in enumerate(lines):
        if _re.search(r'(?i)Pre\s*Bid\s+Detail', line):
            header_idx = i
            break
    if header_idx is not None:
        for data_line in lines[header_idx + 1: header_idx + 12]:
            raw = _re.sub(r'^\[TABLE\]\s*', '', data_line.strip())
            m = DT_PAT.match(raw)
            if m:
                dt = m.group(1).strip()
                venue = raw[m.end():].strip().lstrip('|').strip()
                return dt, venue
    for i, line in enumerate(lines):
        if _re.search(r'(?i)Pre[\s\-]*Bid\s*(Date|Detail|Meeting)', line):
            for data_line in lines[i: i + 20]:
                raw = data_line.strip()
                m = DT_PAT.search(raw)
                if m:
                    return m.group(1).strip(), raw[m.end():].strip().lstrip('|').strip()
            break
    m_broad = _re.search(
        r'(?i)Pre\s*[-\s]?Bid.{0,800}?' + DT_PAT.pattern + r'[ \t]*(.*)',
        doc_text, _re.DOTALL)
    if m_broad:
        dt = m_broad.group(1).strip()
        venue = m_broad.group(2).strip().split('\n')[0].strip().lstrip('|').strip()
        return dt, venue
    return None, None


# ── Consignee / address extractors ──────────────────────────────────────────
_LOC_STOPWORDS = frozenset([
    'india', 'government', 'ministry', 'department', 'office', 'portal',
    'national', 'central', 'public', 'procurement', 'tender', 'limited',
    'ltd', 'pvt', 'corporation', 'authority', 'board', 'council', 'committee',
    'commission', 'agency', 'society', 'trust', 'undertaking', 'services',
    'solutions', 'technologies', 'systems', 'enterprises', 'works', 'supply',
    'project', 'contract', 'bid', 'item', 'description', 'specification',
    'scope', 'quantity', 'unit', 'rate', 'amount', 'total', 'date', 'time',
    'period', 'refer', 'see', 'attached', 'annexure', 'document', 'corrigendum',
    'reporting', 'officer', 'consignee', 'delivery', 'days', 'address',
    'the', 'and', 'for', 'with', 'not', 'nil', 'as', 'per',
])


def _parse_addr_cell_to_city(cell: str):
    cell = _clean_cid(cell).strip()
    cell = _re.sub(r'^\*+', '', cell).strip()
    if not cell or len(cell) < 3:
        return None
    if _re.match(r'^[A-Za-z][A-Za-z\s\-]{2,30}$', cell):
        city = cell.strip()
        if city.lower() not in _LOC_STOPWORDS and _is_valid_place(city):
            return city.title()
    loc, _ = parse_location_from_address(cell)
    if loc:
        return loc
    m = _re.search(r'([A-Za-z][A-Za-z\s]{2,20})[,\-\s]+\d{6}', cell)
    if m:
        candidate = m.group(1).strip().rstrip(',').strip()
        if _is_valid_place(candidate):
            return candidate.title()
    return None


def _extract_consignee_table_addresses(doc_text):
    lines = doc_text.split('\n')
    results, seen_cities = [], []
    addr_col_idx = qty_col_idx = None
    header_found = False
    for line in lines:
        if not line.startswith('[TABLE]'):
            continue
        body   = line[7:].strip()
        cols   = [c.strip() for c in body.split('|')]
        body_c = _clean_cid(body)
        has_addr = bool(_re.search(r'(?i)\bपता\b|\bAddress\b', body_c))
        has_cons = bool(_re.search(r'(?i)Consignee|परेषिती|Reporting', body_c))
        if has_addr and has_cons and not header_found:
            for ci, col in enumerate(cols):
                col_c = _clean_cid(col)
                if (_re.search(r'(?i)\bपता\b|\bAddress\b', col_c)
                        and not _re.search(r'(?i)Delivery|डिलीवरी', col_c)):
                    addr_col_idx = ci
                if _re.search(r'(?i)\bमात्रा\b|\bQty\b|\bQuantity\b', col_c):
                    qty_col_idx = ci
            header_found = True
            continue
        if header_found and addr_col_idx is not None:
            first = _clean_cid(cols[0]).strip() if cols else ''
            if not _re.match(r'^\d+\.?$', first):
                if len(cols) <= 2 and not _re.search(r'(?i)Consignee|Address', body_c):
                    header_found = False
                    addr_col_idx = None
                continue
            raw_addr = cols[addr_col_idx].strip() if addr_col_idx < len(cols) else ''
            qty = (cols[qty_col_idx].strip()
                   if qty_col_idx is not None and qty_col_idx < len(cols) else None)
            city = _parse_addr_cell_to_city(raw_addr)
            if city and city not in seen_cities:
                seen_cities.append(city)
                results.append((city, qty))
    return results


def extract_raw_address(doc_text):
    consignee_addrs = _extract_consignee_table_addresses(doc_text)
    if consignee_addrs:
        return ', '.join(city for city, _ in consignee_addrs)
    m = _re.search(
        r'(?i)\[TABLE\][^\n]*(?:पता\s*/?\s*Address|Address\s*/?\s*पता'
        r'|Consignee\s+Address|Delivery\s+Address|Place\s+of\s+Work'
        r'|Place\s+of\s+Supply|Work\s+Location|Site\s+Location'
        r'|Installation\s+Address|\bAddress\b)\s*\|([^|\n]{5,200})',
        doc_text
    )
    if m:
        val = _clean_cid(m.group(1)).strip()
        if (val and len(val) >= 5 and val.lower() not in _STR_NULL
                and not _re.match(r'^[\u0900-\u097F\s]{1,20}$', val)
                and not _re.search(r'(?i)^(?:मात्रा|quantity|delivery|qty)', val)):
            city = _parse_addr_cell_to_city(val)
            if city:
                return city
    m = _re.search(
        r'(?im)^[ \t]*(?:Delivery\s+Address|Consignee\s+Address|Address)\s*[:\-]\s*(.+?)(?:\n[ \t]*\n|\Z)',
        doc_text, _re.DOTALL
    )
    if m:
        val = _re.sub(r'\s*\n\s*', ', ', m.group(1).strip())
        val = _clean_cid(val).strip()
        if val and len(val) >= 5 and val.lower() not in _STR_NULL:
            city = _parse_addr_cell_to_city(val)
            if city:
                return city
    return None


def extract_location_from_atc(doc_text):
    patterns = [
        r'(?i)city\s+of\s+consignee[^(]*\(([A-Za-z][A-Za-z\s]{2,25})\)',
        r"(?i)consignee['\s]*\s+location[^(]*\(([A-Za-z][A-Za-z\s]{2,25})\)",
        r'(?i)service\s+cent(?:re|er)[^(]*\(([A-Za-z][A-Za-z\s]{2,25})\)',
        r'(?i)delivery\s+at[^(]*\(([A-Za-z][A-Za-z\s]{2,25})\)',
    ]
    for pat in patterns:
        m = _re.search(pat, doc_text)
        if m:
            val = m.group(1).strip()
            if _is_valid_place(val):
                return val
    return None


def _venue_to_city(venue):
    if not venue:
        return None
    pin_m = _re.search(r'([A-Z][a-zA-Z\s]{2,20})[,\-\s]+(\d{6})', venue)
    if pin_m:
        city = pin_m.group(1).strip().rstrip(',').strip()
        if len(city) >= 3 and _is_valid_place(city):
            return city
    _SKIP = _re.compile(
        r'^(?:level|block|sec|sector|plot|floor|room|wing|east|west|north|south'
        r'|building|hall|office|complex|tower|road|street|lane|marg|nagar|phase'
        r'|crpf|cisf|bsf|itbp|ssb|esic|cpwd|\d+)$', _re.I)
    parts = [p.strip() for p in venue.split(',')]
    for part in reversed(parts):
        part = _re.sub(r'[\-\s]*\d{6}\s*$', '', part.strip()).strip()
        if not part:
            continue
        tokens = part.split()
        if all(_SKIP.match(t) for t in tokens):
            continue
        if len(part) >= 3 and _re.search(r'[A-Za-z]{3}', part) and _is_valid_place(part):
            return part
    return None


# ── Tender ID ────────────────────────────────────────────────────────────────
def extract_tender_id_from_text(doc_text):
    head = doc_text[:4000]
    patterns = [
        r'GEM[/\-]\d{4}[/\-][A-Z0-9]+[/\-]\d+',
        r'(?:NIT|RFQ|RFP|IFB|EOI|BID)[/\-][\w/\-]{4,40}',
        r'\d{4}_[A-Z]{2,10}_\d{4,10}(?:_\d+)?',
        r'[A-Z]{2,10}(?:/[A-Z0-9\-]{1,20}){2,5}',
        r'(?i)(?:tender\s*(?:no\.?|id|number)|nit\s*(?:no\.?|number)'
        r'|ref(?:erence)?\s*no\.?|bid\s*(?:no\.?|number)'
        r'|e-?tender\s*no\.?|enquiry\s*no\.?|file\s*no\.?'
        r'|gem\s*bid\s*(?:no\.?|id)|procurement\s*no\.?)'
        r'\s*[:\-]?\s*([A-Z0-9][A-Z0-9/\-_.]{4,60})',
    ]
    for pat in patterns:
        m = _re.search(pat, head)
        if m:
            try:
                val = m.group(1).strip()
            except IndexError:
                val = m.group(0).strip()
            if _re.search(r'\d', val) and len(val) >= 5:
                return val
    return None


# ── Location from text ───────────────────────────────────────────────────────
_ENG_PLACE = r'[A-Za-z][A-Za-z\s\-\.]{1,40}'
_LOC_PATTERNS = [
    (r'(?im)\[TABLE\]\s*(?:place\s*of\s*work|place\s*of\s*supply|work\s*location'
     r'|site\s*location|location\s*of\s*work|delivery\s*(?:address|location)'
     r'|installation\s*site|district|taluka|tehsil|block|mandal|circle'
     r'|zone|region|division|sub.?division|location|place|site|city|town|area)'
     r'\s*[\|:]\s*(' + _ENG_PLACE + r')(?:\s*[\|]|$)'),
    (r'(?im)^[ \t]*(?:place\s*of\s*work|place\s*of\s*supply|work\s*location'
     r'|site\s*location|location\s*of\s*work|delivery\s*address'
     r'|installation\s*site|district|taluka|tehsil|block|mandal|circle'
     r'|location|place|site|city|town)'
     r'\s*[:\-]\s*(' + _ENG_PLACE + r')[ \t]*(?:\n|$)'),
    r'(?m)^\s*\d+\.?\s+\*{3,}[\w\*\s]*?([A-Z][a-zA-Z]{2,25})(?:\s+\d)',
    r'(?i)\[TABLE\][^\n]*(?:पता\s*/\s*Address|address\s*/\s*पता)\s*[\|:]\s*(' + _ENG_PLACE + r')(?:\s*[\|]|$)',
    r'(?m)\[TABLE\][^\n]*[\u0900-\u097F][^\n]*\|\s*([A-Z][a-z]{2,20})\s*(?:\||\d|\n|$)',
    r'(' + _ENG_PLACE + r')\s*[-,]?\s*\b\d{6}\b',
]
_LOC_STAR_PATTERN    = r'\*{3,}([A-Z][a-zA-Z]{2,}(?:\s+[A-Z][a-zA-Z]+)*)(?=\s*[|\n]|\s+\d|$)'
_LOC_TABLE_ADDR_PATTERN = r'(?m)\[TABLE\][^\n]*\|\s*\*{3,}([A-Z][a-zA-Z]{2,}(?:\s+[A-Z][a-zA-Z]+)*)\s*(?:\||\d|$)'
_LOC_TABLE_FULLADDR_PATTERN = (
    r'(?m)\[TABLE\][^\n]*\|([^|\n]{15,}'
    r'(?:DISTT?|\bPIN[-:\s]*\d{6}|'
    r'(?<![A-Z])(?:CG|UP|MP|MH|KA|TN|AP|TS|RJ|GJ|PB|HR|HP|UK|JH|OD|AS|BR|WB|JK|DL)(?![A-Z]))'
    r'[^|\n]*)(?:\||$)'
)


def extract_location_from_text(doc_text: str):
    for m in _re.finditer(_LOC_TABLE_FULLADDR_PATTERN, doc_text):
        raw_addr = m.group(1).strip()
        loc, _st = parse_location_from_address(raw_addr)
        if loc:
            return loc
    _seen_places = []
    for m in _re.finditer(_LOC_TABLE_ADDR_PATTERN, doc_text):
        val = m.group(1).strip()
        if (len(val) >= 3 and val[0].isupper()
                and val.lower() not in _LOC_STOPWORDS
                and val not in _seen_places):
            _seen_places.append(val)
    if _seen_places:
        return ', '.join(_seen_places)
    _seen_star = []
    for m in _re.finditer(_LOC_STAR_PATTERN, doc_text):
        val = m.group(1).strip()
        if (len(val) >= 3 and val[0].isupper()
                and val.lower() not in _LOC_STOPWORDS
                and val not in _seen_star):
            _seen_star.append(val)
    if _seen_star:
        return ', '.join(_seen_star)
    scored = []
    for priority, pat in enumerate(_LOC_PATTERNS):
        for m in _re.finditer(pat, doc_text):
            try:
                val = m.group(1)
            except IndexError:
                continue
            val = val.strip().strip('.,;:/-').strip()
            val = _re.sub(r'\s+', ' ', val)
            if len(val) < 3 or val.isdigit() or not val[0].isupper():
                continue
            low_words = set(val.lower().split())
            if low_words.issubset(_LOC_STOPWORDS) or all(w in _LOC_STOPWORDS for w in low_words):
                continue
            scored.append(((priority * 10) + len(val.split()), val))
    if not scored:
        return None
    scored.sort(key=lambda x: x[0])
    return scored[0][1]


# ── Quantity extractor ───────────────────────────────────────────────────────
_QTY_PATTERNS = [
    r'(?m)^\s*\d+\.?\s+\*{3,}[\w\*\s]*?[A-Z][a-zA-Z]{2,25}\s+(\d+)\s+\d',
    r'(?i)(?:quantity|qty|nos?\.?|number\s+of\s+(?:units?|items?|sets?))\s*[:\-]\s*(\d+)',
    r'(?im)^\[TABLE\][^\n]*(?:qty|quantity|nos?\.?)\s*[\|:]\s*(\d+)',
]


def extract_qty_from_text(doc_text: str, ai_qty):
    if ai_qty is not None:
        try:
            v = int(ai_qty)
            if v > 0:
                return v
        except (TypeError, ValueError):
            pass
    consignee_matches = _re.findall(_QTY_PATTERNS[0], doc_text)
    if consignee_matches:
        try:
            total = sum(int(q) for q in consignee_matches)
            if total > 0:
                return total
        except (ValueError, TypeError):
            pass
    for pat in _QTY_PATTERNS[1:]:
        m = _re.search(pat, doc_text)
        if m:
            try:
                val = int(m.group(1))
                if 1 <= val <= 100000:
                    return val
            except (ValueError, IndexError):
                continue
    return None


# ── Text extractors (PDF / DOCX / TXT) ──────────────────────────────────────
def extract_text_from_pdf(fp):
    text = ''
    with pdfplumber.open(fp) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += _clean_cid(t) + '\n'
            tables = page.extract_tables()
            for table in (tables or []):
                for row in (table or []):
                    if row:
                        cells = [
                            _clean_cid(str(cell).replace('\n', ' ').replace('\r', ' ').strip())
                            for cell in row if cell
                        ]
                        if cells:
                            text += f'[TABLE] {" | ".join(cells)}\n'
    return text


def extract_text_from_docx(fp):
    d = _docx.Document(fp)
    text = ''.join(p.text + '\n' for p in d.paragraphs)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.replace('\n', ' ').replace('\r', ' ').strip() for c in row.cells]
            text += '[TABLE] ' + ' | '.join(cells) + '\n'
    return text


def extract_text_from_txt(fp):
    with open(fp, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


# ── JSON cleaner ─────────────────────────────────────────────────────────────
def clean_json(text):
    text = text.strip()
    if '```' in text:
        for part in text.split('```'):
            part = part.strip()
            if part.lower().startswith('json'):
                part = part[4:].strip()
            if part.startswith('{'):
                text = part
                break
    s = text.find('{')
    if s == -1:
        return text
    depth, in_str, esc, end = 0, False, False, s
    for i, ch in enumerate(text[s:], start=s):
        if esc:
            esc = False
            continue
        if ch == '\\':
            esc = True
            continue
        if ch == '"':
            in_str = not in_str
        if not in_str:
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    end = i
                    break
    candidate = text[s:end + 1]
    try:
        json.loads(candidate)
        return candidate
    except json.JSONDecodeError:
        pass
    try:
        fixed = _re.sub(r',\s*$', '', candidate.rstrip())
        fixed += ']' * (fixed.count('[') - fixed.count(']'))
        fixed += '}' * (fixed.count('{') - fixed.count('}'))
        json.loads(fixed)
        return fixed
    except Exception:
        pass
    return candidate


# ── LLM prompt ───────────────────────────────────────────────────────────────
EXTRACTION_PROMPT = """<start_of_turn>user
You are an expert data extractor for Indian government tender documents (GeM bids, NITs, RFQs, defence procurement).

Read the document below and extract every field. Return ONLY a valid JSON object — no markdown, no text before or after.

## FIELD TYPES

| Type | Format | Fields |
|------|--------|--------|
| BOOLEAN | true or false | prebid_mandatory |
| INTEGER | plain number | qty, ra_no |
| DECIMAL | quoted "0.00" (strip ₹/commas) | estimated_value, tender_budget, emd_fee, epbg_fee, tender_fee, emd_percentage |
| DATE | "YYYY-MM-DD" | published_date, corrigendum_date, date_of_submission |
| DATETIME | "YYYY-MM-DDTHH:MM:SS+05:30" | submission_deadline, prebid_datetime, edited_datetime |
| YES/NO | "Yes" or "No" — never null | emd_exemption, tender_fee_exemption, corrigendum, ra_enabled |
| STRING | quoted or null | all others |

**GeM date format** "DD-MM-YYYY HH:MM:SS" → convert to ISO. Example: "10-03-2026 19:00:00" → "2026-03-10T19:00:00+05:30"

## KEY FIELD RULES

**tender_id** — Unique bid/reference number.
- GeM: "बिड संख्या/Bid Number" → format GEM/YYYY/B/NNNNNNN
- Non-GeM: Tender No, NIT No, Ref No, e-Tender No, Enquiry No. Extract EXACT alphanumeric string.

**title** — GeM: "वस्तु श्रेणी/Item Category" row. Non-GeM: tender subject / Name of Work.

**tender_authority** — Full authority chain.
- GeM: combine "Organisation Name + Department Name + Ministry" → e.g. "Indian Army, Department Of Military Affairs, Ministry Of Defence"
- Non-GeM: Issuing authority name.

**category** — MUST be exactly "Product" or "Service".

**item_description** — Full item description.

**published_date** — DATE only (YYYY-MM-DD).

**submission_deadline** — DATETIME format.

**date_of_submission** — Same date as submission_deadline but DATE only (YYYY-MM-DD).

**qty** — Total units to be procured.

**emd_fee** — EMD rupee amount ONLY.

**emd_percentage** — GeM ePBG Percentage(%) row value.

**estimated_value** — Tender/BOQ value stated directly.

**ra_enabled** — "Yes" if Bid to RA enabled.

**prebid_datetime** — DATE and TIME only from Pre Bid Detail section.

**prebid_mandatory** — true if Pre Bid Date/Time present.

**turnover_criteria** — Return null (extracted server-side).

**evaluation_criteria** — ALL required documents comma-separated.

**location** — Most specific place name. NEVER return "Yes", "No", or boolean.

**state** — Full Indian state name. Never abbreviate.

**corrigendum** — "Yes" only if corrigendum actually issued.

## OUTPUT

Return EXACTLY this JSON structure (null for any field not found):

{"tender_id":null,"title":null,"tender_authority":null,"category":null,"item_description":null,"published_date":null,"submission_deadline":null,"date_of_submission":null,"qty":null,"ra_no":null,"state":null,"location":null,"evaluation_criteria":null,"estimated_value":null,"tender_budget":null,"turnover_criteria":null,"emd_fee":null,"epbg_fee":null,"emd_percentage":null,"emd_exemption":"No","tender_fee":null,"tender_fee_exemption":"No","prebid_mandatory":false,"prebid_datetime":null,"ra_enabled":"No","corrigendum":"No","corrigendum_date":null,"documents_link":null,"attachments":null,"attachments_name":null,"remarks":null,"edited_datetime":null}
<end_of_turn>
<start_of_turn>model
"""


# ── LLM helpers ──────────────────────────────────────────────────────────────
def get_loaded_model(base_url: str, fallback: str) -> str:
    try:
        resp = req_lib.get(f'{base_url}/models', timeout=5)
        resp.raise_for_status()
        models = resp.json().get('data', [])
        if models:
            return models[0]['id']
        return fallback
    except Exception as e:
        print(f'[model-detect] Cannot reach LMStudio: {e}')
        return fallback
